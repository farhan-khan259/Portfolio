from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins (narrow to maximise space) ─────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin   = Cm(1.6)
    section.right_margin  = Cm(1.6)

# ── Colour palette ───────────────────────────────────────────────────────────
DARK   = RGBColor(0x1A, 0x1A, 0x2E)   # near-black
ACCENT = RGBColor(0x2D, 0x6A, 0xCE)   # professional blue
GRAY   = RGBColor(0x55, 0x55, 0x66)   # muted gray
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: set paragraph spacing ────────────────────────────────────────────
def spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def set_font(run, size, bold=False, italic=False, color=None):
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.name  = "Calibri"
    if color:
        run.font.color.rgb = color

# ── Helper: horizontal rule ──────────────────────────────────────────────────
def add_rule(paragraph, color=ACCENT, thickness=8):
    """Add a bottom border to a paragraph (acts as a section divider)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(thickness))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Helper: section heading ───────────────────────────────────────────────────
def add_section(title):
    p = doc.add_paragraph()
    spacing(p, before=7, after=1)
    run = p.add_run(title.upper())
    set_font(run, 9.5, bold=True, color=ACCENT)
    add_rule(p, ACCENT, 6)
    return p

# ── Helper: bullet ────────────────────────────────────────────────────────────
def add_bullet(text, indent=0.35):
    p = doc.add_paragraph(style="List Bullet")
    spacing(p, before=0, after=1, line=10.5)
    p.paragraph_format.left_indent  = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    run = p.add_run(text)
    set_font(run, 9.2, color=DARK)
    return p

# ── Helper: key-value inline (for experience header) ─────────────────────────
def add_exp_header(title, company, period):
    p = doc.add_paragraph()
    spacing(p, before=5, after=1)
    r1 = p.add_run(f"{title}  ·  ")
    set_font(r1, 9.5, bold=True, color=DARK)
    r2 = p.add_run(company)
    set_font(r2, 9.5, bold=True, color=ACCENT)
    r3 = p.add_run(f"    {period}")
    set_font(r3, 8.8, italic=True, color=GRAY)
    return p

# ── Helper: project header ────────────────────────────────────────────────────
def add_project_header(name, stack):
    p = doc.add_paragraph()
    spacing(p, before=5, after=1)
    r1 = p.add_run(name)
    set_font(r1, 9.5, bold=True, color=DARK)
    r2 = p.add_run(f"  —  {stack}")
    set_font(r2, 8.6, italic=True, color=GRAY)
    return p

# ════════════════════════════════════════════════════════════════════════════════
# NAME & CONTACT HEADER
# ════════════════════════════════════════════════════════════════════════════════
name_p = doc.add_paragraph()
spacing(name_p, before=0, after=2)
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = name_p.add_run("Ahmed Usman")
r.font.name = "Calibri"
r.font.size = Pt(24)
r.font.bold = True
r.font.color.rgb = DARK

title_p = doc.add_paragraph()
spacing(title_p, before=0, after=2)
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rt = title_p.add_run("AI/ML Engineer  ·  LLM Applications  ·  RAG  ·  Voice Agents  ·  MLOps")
set_font(rt, 9.8, italic=False, color=ACCENT)

contact_p = doc.add_paragraph()
spacing(contact_p, before=1, after=4)
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = contact_p.add_run(
    "+92 335 0707006   |   ahmadusman050@gmail.com   |   "
    "github.com/ahmedosm0   |   linkedin.com/in/ahmedusman050"
)
set_font(rc, 8.8, color=GRAY)

# ════════════════════════════════════════════════════════════════════════════════
# PROFESSIONAL SUMMARY
# ════════════════════════════════════════════════════════════════════════════════
add_section("Professional Summary")
summary_p = doc.add_paragraph()
spacing(summary_p, before=2, after=2, line=11)
rs = summary_p.add_run(
    "AI/ML Engineer with a strong track record of shipping production LLM applications, "
    "RAG pipelines, and agentic workflows. Experienced across the full stack — FastAPI, "
    "Next.js, Supabase/PostgreSQL, LangGraph, and cloud deployment — with hands-on work "
    "in voice automation, computer vision, and multi-provider model orchestration. "
    "Focused on systems that are reliable, cost-efficient, and built to last in production."
)
set_font(rs, 9.2, color=DARK)

# ════════════════════════════════════════════════════════════════════════════════
# TECHNICAL SKILLS  (two-column table for compactness)
# ════════════════════════════════════════════════════════════════════════════════
add_section("Technical Skills")

skills = [
    ("Languages",         "Python, TypeScript, JavaScript, C++"),
    ("LLM & Agents",      "LangChain, LangGraph, CrewAI, RAG, vector search, tool calling, semantic caching"),
    ("LLM Providers",     "OpenAI, Anthropic, Mistral, Groq, Qwen, OpenRouter, Ollama, Hugging Face"),
    ("AI/ML",             "TensorFlow, Scikit-learn, YOLOv8, U-Net, BiLSTM, OpenCV, DeepFace"),
    ("Voice & Speech",    "Retell AI, ElevenLabs, Whisper, Coqui TTS, Librosa"),
    ("Backend & APIs",    "FastAPI, Django, Node.js, Express.js, WebSockets, SSE, REST"),
    ("Frontend",          "Next.js, React, React Native, Expo, TanStack Query, Zustand"),
    ("Databases",         "PostgreSQL, Supabase, MongoDB, Redis, Pinecone, ChromaDB, Cloudflare R2"),
    ("Automation",        "n8n, Shopify GraphQL, Twilio, Stripe, Buffer, Gmail API, Klaviyo, GA4, PostHog"),
    ("MLOps & DevOps",    "Docker, Jenkins, CI/CD, Sentry, Railway, AWS, model monitoring"),
]

tbl = doc.add_table(rows=len(skills), cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, (label, value) in enumerate(skills):
    cells = tbl.rows[i].cells
    # left cell — label
    lp = cells[0].paragraphs[0]
    spacing(lp, before=1, after=1)
    lr = lp.add_run(label)
    set_font(lr, 8.8, bold=True, color=DARK)
    cells[0].width = Inches(1.3)
    # right cell — value
    rp = cells[1].paragraphs[0]
    spacing(rp, before=1, after=1)
    rr = rp.add_run(value)
    set_font(rr, 8.8, color=DARK)
    # remove borders for a clean look
    for cell in cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top","left","bottom","right","insideH","insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBorders.append(el)
        tcPr.append(tcBorders)

# ════════════════════════════════════════════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ════════════════════════════════════════════════════════════════════════════════
add_section("Professional Experience")

add_exp_header("Associate AI Engineer", "Tech Emulsion  (Full-time)", "Jul 2025 – Present")
add_bullet("Design and deploy LLM-powered conversational agents, RAG pipelines, and automation workflows for production AI products.")
add_bullet("Build backend services and integrations using Python, Django, FastAPI, n8n, AWS, and third-party APIs.")

add_exp_header("AI/ML Engineer", "DevK System  (Part-time)", "Aug 2024 – Jun 2025")
add_bullet("Built LLM applications including RAG systems, autonomous agents, and agentic workflow prototypes.")
add_bullet("Applied MLOps practices: Dockerized deployment, CI/CD workflows, model monitoring, and latency-focused design.")

add_exp_header("Back-End Developer", "Brandora  (Part-time)", "Sep 2022 – Jul 2024")
add_bullet("Developed REST APIs with Node.js, Express.js, and MongoDB; integrated with React frontends and optimised database queries.")

# ════════════════════════════════════════════════════════════════════════════════
# SELECTED PROJECTS
# ════════════════════════════════════════════════════════════════════════════════
add_section("Selected Projects")

# 1 ── The Meatery
add_project_header(
    "The Meatery — E-Commerce Intelligence & AI Voice-Agent Platform",
    "Next.js 14, TypeScript, Retell AI, n8n, Shopify, OpenAI, Anthropic, Twilio"
)
add_bullet("Built a revenue-ops and analytics platform for a US premium-meat brand, unifying margin monitoring, inventory velocity, reorder automation, marketing attribution, and AI voice-agent CRM workflows.")
add_bullet("Engineered Retell AI inbound/outbound calling agents for abandoned-cart recovery and win-back campaigns, backed by live Shopify cart lookup, dynamic discount generation, DNC enforcement, and duplicate-call prevention.")
add_bullet("Added nightly n8n pipelines that analyse call transcripts with Claude to generate objection rebuttals and prompt-improvement recommendations, then sync updates into Retell knowledge bases.")

# 2 ── AVL Copilot
add_project_header(
    "AVL Copilot — AI Technical-Support Assistant",
    "FastAPI, LangGraph, Pinecone, Redis, Supabase, Stripe, OpenAI Responses API"
)
add_bullet("Deployed a multimodal RAG + agentic support assistant for field AVL technicians with a strictly sequential LangGraph pipeline: semantic cache → intent detection → RAG → web-search fallback → streamed SSE generation.")
add_bullet("Implemented dual-tier semantic cache, circuit breakers, domain allowlisting, image diagnostics, tiered model routing, Stripe billing, and a live SSE log viewer.")

# 3 ── Good Food Project
add_project_header(
    "Good Food Project — AI Social Content Engine",
    "FastAPI, Supabase, LangGraph, Next.js 15, React 19, TypeScript, Cloudflare R2"
)
add_bullet("Built an AI content-generation platform for a UK organic food brand: LangGraph multi-node pipeline, multi-provider LLM routing, RAG over brand knowledge, Buffer scheduled publishing, and Sentry observability.")
add_bullet("Delivered a canvas-based image editor (react-konva), Zustand state management, TanStack Query data fetching, and Cypress E2E coverage.")

# 4 ── LinkedIn Outreach
add_project_header(
    "LinkedIn Outreach Automation — AI Vision-Driven Connection Agent",
    "FastAPI, React 19, Supabase, Playwright, Qwen-VL, DashScope, SSE"
)
add_bullet("Built a B2B lead-generation platform automating LinkedIn outreach through a real-time four-step pipeline (scrape → enrich → AI note → send) with a three-tier vision decision cascade to handle UI variants.")
add_bullet("Implemented InMail-credit budgeting, model-pool rotation across nine Qwen models, multi-tenant encrypted session storage, and resume-from-send mode.")

# 5 ── Lost-N-Find  (compact)
add_project_header(
    "Lost-N-Find — AI Campus Lost-and-Found System",
    "FastAPI, Supabase, React Native, Expo, WebSockets, Mistral"
)
add_bullet("Hybrid claim-scoring engine combining deterministic matching with LLM semantic scoring; real-time WebSocket notifications, JWT auth, and AI-powered ownership verification.")

# ════════════════════════════════════════════════════════════════════════════════
# EDUCATION & CERTIFICATIONS  (side-by-side table)
# ════════════════════════════════════════════════════════════════════════════════
add_section("Education & Certifications")

edu_tbl = doc.add_table(rows=1, cols=2)
edu_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Education cell
left = edu_tbl.rows[0].cells[0]
lp = left.paragraphs[0]
spacing(lp, before=2, after=1)
lr1 = lp.add_run("University of Engineering and Technology, Peshawar\n")
set_font(lr1, 9.2, bold=True, color=DARK)
lr2 = lp.add_run("B.S. Electrical Computing & Communication Engineering")
set_font(lr2, 8.8, color=GRAY)

# Certifications cell
right = edu_tbl.rows[0].cells[1]
rp = right.paragraphs[0]
spacing(rp, before=2, after=1)
rr1 = rp.add_run("Certifications: ")
set_font(rr1, 9.2, bold=True, color=DARK)
rr2 = rp.add_run(
    "Machine Learning Specialization · Deep Learning Specialization · "
    "ML in Production · TensorFlow for AI/ML/DL · "
    "Generative AI with LangChain & Hugging Face"
)
set_font(rr2, 8.6, color=GRAY)

# Remove table borders
for row in edu_tbl.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for side in ("top","left","bottom","right","insideH","insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "nil")
            tcBorders.append(el)
        tcPr.append(tcBorders)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r"c:\Users\ahmad\Downloads\cv_project\Ahmed_Usman_CV_v2.docx"
doc.save(out)
print(f"Saved: {out}")
