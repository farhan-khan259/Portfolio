# -*- coding: utf-8 -*-
"""
Generates Ahmed_Usman_CV_final.docx — layout matches Ahmed's_CV.pdf:
  Times New Roman, centered header, bold section headers with bottom-border rule,
  inline skills, right-tab-aligned experience dates, categorised projects with
  bold+underlined project titles and 2-3 descriptive bullets each.
Target: 3 pages.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(1.8)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin   = Cm(2.3)
    sec.right_margin  = Cm(2.3)

FONT    = "Times New Roman"
SZ_BODY = 10.5
SZ_NAME = 20
SZ_TITL = 12
SZ_SEC  = 12
BLACK   = RGBColor(0x00, 0x00, 0x00)
BLUE    = RGBColor(0x11, 0x55, 0xCC)

# ── Core helpers ─────────────────────────────────────────────────────────────

def run(para, text, size=SZ_BODY, bold=False, italic=False,
        underline=False, color=BLACK):
    r = para.add_run(text)
    r.font.name      = FONT
    r.font.size      = Pt(size)
    r.font.bold      = bold
    r.font.italic    = italic
    r.font.underline = underline
    r.font.color.rgb = color
    return r

def sp(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line is not None:
        pf.line_spacing = Pt(line)

def bottom_border(para, sz=6):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "000000")
    pBdr.append(bot)
    pPr.append(pBdr)

def section_head(title):
    p = doc.add_paragraph()
    sp(p, before=9, after=3)
    run(p, title, size=SZ_SEC, bold=True)
    bottom_border(p)
    return p

def exp_head(left_text, date_text):
    p = doc.add_paragraph()
    sp(p, before=5, after=1)
    pf = p.paragraph_format
    pf.tab_stops.add_tab_stop(Inches(6.15), WD_TAB_ALIGNMENT.RIGHT)
    run(p, left_text, bold=True)
    run(p, f"\t{date_text}")
    return p

def bullet(parts, indent_in=0.28, hanging_in=0.18):
    p = doc.add_paragraph(style="List Bullet")
    sp(p, before=1, after=1, line=SZ_BODY + 1)
    pf = p.paragraph_format
    pf.left_indent       = Inches(indent_in)
    pf.first_line_indent = Inches(-hanging_in)
    for text, bold, underline, color in parts:
        run(p, text, bold=bold, underline=underline, color=color)
    return p

def skill_row(label, value):
    p = doc.add_paragraph()
    sp(p, before=0, after=1, line=SZ_BODY + 1)
    run(p, label + ": ", bold=True)
    run(p, value)
    return p

def cat_head(title):
    """Category sub-heading inside Projects section."""
    p = doc.add_paragraph()
    sp(p, before=6, after=2)
    run(p, title, bold=True)
    return p

def proj_title(name, stack=""):
    """Bold + underlined project name line with optional stack note."""
    p = doc.add_paragraph()
    sp(p, before=5, after=1)
    run(p, name, bold=True, underline=True, color=BLUE)
    if stack:
        run(p, f"  |  {stack}", italic=True, color=BLACK)
    return p

# ════════════════════════════════════════════════════════════════════════════
# HEADER
# ════════════════════════════════════════════════════════════════════════════
np_ = doc.add_paragraph()
sp(np_, before=0, after=2)
np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(np_, "Ahmed Usman", size=SZ_NAME, bold=True)

tp_ = doc.add_paragraph()
sp(tp_, before=0, after=2)
tp_.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(tp_, "AI/ML Engineer", size=SZ_TITL)

c1 = doc.add_paragraph()
sp(c1, before=0, after=1)
c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(c1, "\U0001F4DE ")
run(c1, "+92 335 0707006", underline=True, color=BLUE)
run(c1, "  |  \U00002709 ")
run(c1, "ahmadusman050@gmail.com", underline=True, color=BLUE)
run(c1, "  |")

c2 = doc.add_paragraph()
sp(c2, before=0, after=4)
c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(c2, "\U0001F310 ")
run(c2, "github.com/ahmedosm0", underline=True, color=BLUE)
run(c2, "  |  \U0001F4BC ")
run(c2, "linkedin.com/in/ahmedusman050", underline=True, color=BLUE)

# ════════════════════════════════════════════════════════════════════════════
# OBJECTIVE
# ════════════════════════════════════════════════════════════════════════════
section_head("About Me")
op = doc.add_paragraph()
sp(op, before=2, after=2, line=SZ_BODY + 2)
run(op, (
    "I am an AI/ML Engineer with a genuine passion for building intelligent systems that solve "
    "real problems in production. Over the past few years I have designed and shipped full-stack "
    "AI products spanning LLM applications, RAG pipelines, agentic workflows, voice automation, "
    "and applied computer vision — working across the entire stack from model orchestration and "
    "retrieval systems to backend APIs and user-facing interfaces. "
    "I thrive in environments where research meets engineering: I enjoy taking an idea from a "
    "rough prototype all the way to a reliable, monitored, cost-efficient deployment. "
    "I work well in cross-functional teams, communicate clearly about technical trade-offs, and "
    "take ownership of the quality of what I ship."
))

# ════════════════════════════════════════════════════════════════════════════
# SKILLS
# ════════════════════════════════════════════════════════════════════════════
section_head("Skills")

for label, value in [
    ("Programming Languages",    "Python, JavaScript, TypeScript, C++"),
    ("Web Frameworks",           "FastAPI, Django, Flask, Node.js, Express.js"),
    ("Frontend & Mobile",        "Next.js, React, React Native, Expo, TanStack Query, Zustand"),
    ("Databases",                "PostgreSQL, Supabase, MongoDB, MySQL, SQLite, Redis, Pinecone, ChromaDB"),
    ("AI/ML",                    "TensorFlow, Scikit-learn, OpenCV, DeepFace, YOLOv8, U-Net, BiLSTM"),
    ("LLM & Agent Systems",      "LangChain, LangGraph, CrewAI, LangSmith, RAG, vector search, semantic caching, tool calling"),
    ("LLM Providers",            "OpenAI, Anthropic, Mistral, Groq, Qwen, DashScope, OpenRouter, Ollama, Hugging Face"),
    ("Speech & Voice",           "Retell AI, ElevenLabs, Whisper, Coqui TTS, Librosa, SciPy"),
    ("Automation & Integrations","n8n, Shopify GraphQL, Twilio, Stripe, Buffer, Gmail API, Google Sheets, Klaviyo, GA4, PostHog"),
    ("MLOps & DevOps",           "Docker, Jenkins, Git, GitHub, CI/CD, Sentry, Railway, AWS, model monitoring"),
    ("Soft Skills",              "Communication, Team Collaboration, Problem-Solving, Fast Learning"),
    ("Languages",                "English (C1), German (A2)"),
]:
    skill_row(label, value)

# ════════════════════════════════════════════════════════════════════════════
# EXPERIENCE
# ════════════════════════════════════════════════════════════════════════════
section_head("Experience")

exp_head("Associate AI Engineer at Tech Emulsion - (full time)", "(Jul, 2025 - Present)")
bullet([("Design and deploy LLM-driven conversational agents, RAG pipelines, and automation "
         "workflows by integrating APIs, structured backends, and third-party tools to deliver "
         "scalable AI solutions.", False, False, BLACK)])
bullet([("Develop and maintain backend services and integrations using Python, Django, FastAPI, "
         "n8n, and AWS, collaborating with cross-functional teams to deliver end-to-end AI products.", False, False, BLACK)])

exp_head("AI/ML Engineer at DevK System - (part time)", "(Aug, 2024 - Jun, 2025)")
bullet([("Designed and developed LLM applications including RAG systems, autonomous agents, "
         "and agentic workflow prototypes covering the full lifecycle from research to deployment.", False, False, BLACK)])
bullet([("Designed and deployed ML models with attention to latency, scalability, and "
         "maintainability; applied MLOps practices using Docker, CI/CD pipelines, and model monitoring.", False, False, BLACK)])

exp_head("Back-End Developer at Brandora - (part time)", "(Sep, 2022 - Jul, 2024)")
bullet([("Developed REST APIs with Node.js, Express.js, and MongoDB; integrated backend services "
         "with React frontends and implemented authentication, query optimisation, and reusable "
         "API modules.", False, False, BLACK)])

# ════════════════════════════════════════════════════════════════════════════
# PROJECTS
# ════════════════════════════════════════════════════════════════════════════
section_head("Projects")

# ── Category 1: Full-Stack AI Platforms ─────────────────────────────────────
cat_head("Full-Stack AI Platforms")

# 1. The Meatery
proj_title("The Meatery – E-Commerce Intelligence & AI Voice-Agent Platform",
           "Next.js 14, TypeScript, Retell AI, n8n, Shopify, OpenAI, Anthropic, Twilio")
bullet([("Built a full-stack revenue-operations and analytics platform for a US premium-meat "
         "e-commerce brand, unifying margin/COGS monitoring, inventory velocity tracking, "
         "reorder automation, marketing attribution (Klaviyo, GA4, Search Console, PostHog), "
         "competitor price scraping, and GrowthBook A/B pricing experiments within a single "
         "operator dashboard.", False, False, BLACK)])
bullet([("Engineered Retell AI inbound and outbound calling agents for abandoned-cart recovery, "
         "win-back campaigns, prospecting, and post-delivery support — backed by an Express.js "
         "callback server that performs live Shopify cart lookup, generates dynamic discount codes, "
         "delivers SMS via Twilio, enforces DNC lists, applies per-lead cooldowns, and prevents "
         "duplicate calls.", False, False, BLACK)])
bullet([("Designed nightly n8n pipelines that analyse call transcripts with Claude to produce "
         "objection rebuttals, competitive battle cards, and prompt-improvement suggestions, then "
         "automatically sync the updates into agent-specific Retell knowledge bases; added a "
         "streaming LLM concierge shopping agent and AI-driven product-content generation.", False, False, BLACK)])

# 2. AVL Copilot
proj_title("AVL Copilot – AI Technical-Support Copilot",
           "FastAPI, LangGraph, Pinecone, Redis, Supabase, Stripe, OpenAI Responses API")
bullet([("Built and deployed a multimodal RAG and agentic support assistant that helps field "
         "technicians troubleshoot Audio, Video, and Lighting (AVL) equipment, retrieve "
         "manufacturer manuals, and answer complex engineering questions in real time via "
         "FastAPI with server-sent event (SSE) streaming.", False, False, BLACK)])
bullet([("Orchestrated a strictly sequential LangGraph pipeline — semantic cache check → "
         "conversation summary → intent detection → Pinecone RAG retrieval → query routing → "
         "Serper/ScraperAPI web-search fallback with Jina Reader content extraction → streamed "
         "generation — enforced with circuit breakers, domain allowlisting, and "
         "manufacturer-support reranking.", False, False, BLACK)])
bullet([("Added dual-tier semantic cache (Redis), image-based diagnostics, dynamic token "
         "budgeting, tiered model routing, Supabase auth with conversation checkpointing, Stripe "
         "billing with per-user quotas, a PDF/URL manual ingestion pipeline, metrics endpoints, "
         "and a live SSE log viewer for real-time operational observability.", False, False, BLACK)])

# 3. Good Food Project
proj_title("Good Food Project – AI Social Content Engine",
           "FastAPI, Supabase, LangGraph, Next.js 15, React 19, TypeScript, Cloudflare R2")
bullet([("Built an AI-powered content generation platform for a UK organic food brand, using a "
         "LangGraph multi-node pipeline to create brand-aligned social media posts and route them "
         "through structured human review and approval workflows before scheduled publishing.", False, False, BLACK)])
bullet([("Implemented a FastAPI and Supabase backend with async background jobs, multi-provider "
         "LLM routing with rate-limit handling and token-cost tracking, RAG over a brand knowledge "
         "base (ChromaDB), computer-vision image matching, automated quote-image composition, "
         "Buffer API and APScheduler for scheduled publishing, Cloudflare R2 for asset storage, "
         "and Sentry for observability.", False, False, BLACK)])
bullet([("Delivered a canvas-based image editor built with react-konva, Zustand for state "
         "management, TanStack Query for data fetching, Supabase SSR authentication, and "
         "Cypress end-to-end test coverage on the frontend.", False, False, BLACK)])

# 4. LinkedIn Outreach Automation
proj_title("LinkedIn Outreach Automation – AI Vision-Driven Connection Agent",
           "FastAPI, React 19, Supabase, Playwright, Qwen-VL, DashScope, SSE")
bullet([("Built a B2B lead-generation platform that automates LinkedIn outreach through a "
         "real-time four-step pipeline — scrape profile → enrich data → generate AI personalised "
         "connection note → send connection request — with live progress updates streamed to the "
         "React frontend via SSE.", False, False, BLACK)])
bullet([("Designed a three-tier AI decision cascade using scoped action-bar snapshots, Qwen-VL "
         "visual element screenshots, and deterministic label scanning to reliably handle "
         "LinkedIn UI variants and modal state changes; built model-pool rotation across nine "
         "Qwen models with automatic rate-limit retry and InMail-credit budget tracking.", False, False, BLACK)])
bullet([("Implemented multi-tenant session handling with AES-encrypted Supabase credential "
         "storage, concurrent per-user pipeline slots, auth recovery, idempotent cancellation, "
         "resume-from-send mode to restart interrupted campaigns, and a connection analytics "
         "dashboard.", False, False, BLACK)])

# ── Category 2: Mobile Applications ─────────────────────────────────────────
cat_head("Mobile Applications")

# 5. Lost-N-Find
proj_title("Lost-N-Find – AI Campus Lost-and-Found Mobile App",
           "React Native, Expo, FastAPI, Supabase, WebSockets, Mistral")
bullet([("Designed and built a cross-platform mobile application (React Native + Expo) for a "
         "university campus that digitises the entire lost-and-found process — students report "
         "lost or found items, an admin reviews AI-assisted ownership claims, and approved "
         "claimants are notified and connected through the app.", False, False, BLACK)])
bullet([("Developed a hybrid claim-scoring engine that combines deterministic field matching "
         "(exact/fuzzy string comparison on item attributes) with LLM semantic scoring via "
         "Mistral to correctly handle typos, synonyms, and numerical variants, making "
         "verification robust without relying solely on the language model.", False, False, BLACK)])
bullet([("Implemented lost-to-found item auto-matching, category-based location prediction, "
         "JWT authentication, real-time WebSocket push notifications for claim status updates, "
         "and graceful AI fallback behaviour that keeps the app functional when the model "
         "is unavailable.", False, False, BLACK)])

# ── Category 3: Voice Agents & Speech Processing ─────────────────────────────
cat_head("Voice Agents & Speech Processing")

# 6. AI Health Receptionist
proj_title("AI Health Receptionist Voice Agent",
           "ElevenLabs, Dentally, Pabau, Payment APIs")
bullet([("Developed a healthcare voice agent using ElevenLabs TTS and STT that autonomously "
         "handles appointment booking, cancellation, and rescheduling through natural spoken "
         "conversation — eliminating routine front-desk calls without requiring any staff "
         "involvement.", False, False, BLACK)])
bullet([("Integrated directly with practice management systems (Dentally, Pabau) for real-time "
         "calendar slot lookup and appointment updates, and connected a payment processing "
         "workflow to collect deposits and copayments within the same call flow.", False, False, BLACK)])

# 7. Speech Enhancement (FYP)
proj_title("Speech Enhancement – Final Year Project",
           "U-Net, BiLSTM, PyTorch, Voice Bank, DEMAND dataset, Gradio")
bullet([("Designed and trained a deep learning speech enhancement model using a hybrid "
         "U-Net + BiLSTM architecture on the Voice Bank and DEMAND datasets, targeting "
         "realistic noisy speech across a wide range of acoustic environments and noise types.", False, False, BLACK)])
bullet([("Achieved measurable objective speech-quality improvements over the noisy baseline: "
         "PESQ +8.45%, STOI +1.56%, ESTOI +12.05%; delivered a Gradio web interface for "
         "real-time audio upload, processing, and playback, and published the research output "
         "documenting the architecture and evaluation results.", False, False, BLACK)])

# ════════════════════════════════════════════════════════════════════════════
# EDUCATION
# ════════════════════════════════════════════════════════════════════════════
section_head("Education")

ep = doc.add_paragraph()
sp(ep, before=3, after=1)
run(ep, "University of Engineering and Technology, Peshawar\n", bold=True)
run(ep, "B.S. Electrical Computing & Communication Engineering")

# ════════════════════════════════════════════════════════════════════════════
# CERTIFICATIONS
# ════════════════════════════════════════════════════════════════════════════
section_head("Certifications")

for cert in [
    "Machine Learning Specialization",
    "Deep Learning Specialization",
    "Machine Learning in Production",
    "TensorFlow for AI, ML and DL",
    "Complete Generative AI Course With LangChain and Hugging Face",
]:
    bullet([(cert, False, True, BLUE)], indent_in=0.28)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r"c:\Users\ahmad\Downloads\cv_project\Ahmed_Usman_CV_final3.docx"
doc.save(out)
print(f"Saved: {out}")
